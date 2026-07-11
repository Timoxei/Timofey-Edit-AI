#!/usr/bin/env python3
"""Convert the Aber Kawas blitz-file markdown into a formatted .docx that
Google Docs can open/convert natively. Handles headings, bold, inline links,
bullet lists, blockquotes, tables, and horizontal rules."""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\Users\user\Timofey-Edit-AI\kawas_research\Aber-Kawas-Blitz-File.md"
OUT = r"C:\Users\user\Timofey-Edit-AI\kawas_research\Aber-Kawas-Blitz-File.docx"

LINK_RE = re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
BAREURL_RE = re.compile(r"(?<![\(\]])\b(https?://[^\s|]+)")


def add_hyperlink(paragraph, url, text, bold=False):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    if bold:
        rPr.append(OxmlElement("w:b"))
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_formatted_runs(paragraph, text):
    """Parse a text span for [label](url), bare urls, and **bold**; add runs."""
    # Tokenize by markdown links first.
    pos = 0
    tokens = []  # (kind, a, b)
    for m in LINK_RE.finditer(text):
        if m.start() > pos:
            tokens.append(("text", text[pos:m.start()], None))
        tokens.append(("link", m.group(1), m.group(2)))
        pos = m.end()
    if pos < len(text):
        tokens.append(("text", text[pos:], None))

    for kind, a, b in tokens:
        if kind == "link":
            add_hyperlink(paragraph, b, a)
            continue
        # Within plain text: handle bare urls, then bold.
        sub = a
        upos = 0
        for um in BAREURL_RE.finditer(sub):
            if um.start() > upos:
                _add_bold_segments(paragraph, sub[upos:um.start()])
            url = um.group(1).rstrip(".,);")
            add_hyperlink(paragraph, url, url)
            upos = um.start() + len(url)
        if upos < len(sub):
            _add_bold_segments(paragraph, sub[upos:])


def _add_bold_segments(paragraph, text):
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        r = paragraph.add_run(m.group(1))
        r.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def main():
    with open(SRC, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Table block
        if stripped.startswith("|") and i + 1 < n and set(lines[i + 1].strip()) <= set("|-: "):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            rows = []
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for j, h in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.paragraphs[0].text = ""
                add_formatted_runs(cell.paragraphs[0], h)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for j, val in enumerate(row[:len(header)]):
                    cells[j].paragraphs[0].text = ""
                    add_formatted_runs(cells[j].paragraphs[0], val)
            doc.add_paragraph()
            continue

        if stripped == "---":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")
            pbdr.append(bottom)
            pPr.append(pbdr)
            i += 1
            continue

        if stripped.startswith("#"):
            m = re.match(r"(#+)\s+(.*)", stripped)
            level = min(len(m.group(1)), 4)
            heading = doc.add_heading(level=level)
            add_formatted_runs(heading, m.group(2))
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            add_formatted_runs(p, stripped.lstrip(">").strip())
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        if stripped == "":
            i += 1
            continue

        p = doc.add_paragraph()
        add_formatted_runs(p, stripped)
        i += 1

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    main()
